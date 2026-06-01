import os
import pypdf
import docx


def extraire_texte_fichier(file_path):
    """
    Extrait le texte brut d'un fichier en fonction de son extension.
    Supporte : PDF, DOCX, TXT, CSV.
    Retourne une chaine de caracteres brute.
    """
    if not os.path.exists(file_path):
        return ""

    ext = os.path.splitext(file_path)[1].lower()
    texte = ""

    try:
        if ext == '.pdf':
            # Extraction PDF
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                pages_text = []
                # Limiter a 50 pages max pour eviter les problemes de performance
                for i in range(min(len(reader.pages), 50)):
                    txt = reader.pages[i].extract_text()
                    if txt:
                        pages_text.append(txt)
                texte = "\n".join(pages_text)

        elif ext in ('.docx', '.doc'):
            # Extraction Word
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            # Extraire aussi des tables si besoin
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            paragraphs.append(cell.text)
            texte = "\n".join(paragraphs)

        elif ext in ('.txt', '.csv'):
            # Extraction fichier texte brut
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                texte = f.read(50000) # Limiter a 50k caracteres

    except Exception as e:
        # Journaliser ou simplement ignorer l'erreur pour ne pas bloquer l'upload
        print(f"Erreur d'extraction de texte pour {file_path} : {e}")
        return ""

    # Nettoyer les espaces multiples
    texte_nettoye = " ".join(texte.split())
    return texte_nettoye


def indexer_document_texte(doc):
    """
    Extrait le texte du fichier du document et l'enregistre dans ses metadonnees.
    """
    if not doc.fichier:
        return
    
    from django.conf import settings
    from .models import MetadataDocument
    
    file_path = os.path.join(settings.MEDIA_ROOT, str(doc.fichier))
    texte = extraire_texte_fichier(file_path)
    
    meta, _ = MetadataDocument.objects.get_or_create(document=doc)
    meta.texte_extrait = texte
    meta.save()
